"""DOCX → 공통 문서 모델(DocumentModel) 백엔드.

PDF 백엔드(pdf_backend.py)와 달리 DOCX 는 OOXML 트리를 python-docx 로 직접 순회할 수
있어 판별/OCR 단계 없이 구조를 그대로 읽어낸다. 핵심 설계:

  - **입력 정규화 선행**: python-docx 는 `w:tr`/`w:tc`/`w:p` 의 직계 자식만 보므로,
    내용 컨트롤(`w:sdt`)·필드(`w:fldSimple`) 껍데기에 든 셀·run 은 아예 안 보인다.
    파싱을 시작하기 전에 `_normalize_content_controls()` 가 그 껍데기만 벗겨낸다
    (그래야 표 그리드 계산까지 원래 의도대로 동작한다).
  - **텍스트박스**: 문단 안에 문단·표가 다시 들어있는 구조(`w:txbxContent`)라 껍데기
    벗기기로는 못 푼다. `_textbox_blocks()` 가 안쪽 문단/표를 본문과 같은 변환 경로로
    재귀 처리해 최상위 블록으로 내보낸다(hwpx 백엔드의 도형 텍스트 처리와 같은 관례).
  - **표**: `_extract_table()` 이 python-docx 의 grid(row.cells, 병합 셀은 앵커 셀
    참조가 반복됨)를 직접 순회해 병합 영역(merges)과 셀 안에 실제 중첩된 표
    (nested_tables, cell.tables 로 재귀)를 정확히 복원한다. Docling 의 nested bool
    근사(§docling_adapter._has_nested_structure)와 달리 실제 구조를 얻는다.
  - **머리말/꼬리말**: Block.section 필드(body/header/footer, 그리고 실제로 켜져 있을
    때만 header_first/footer_first/header_even/footer_even)로 본문과 구분한다.
    Block.type 은 그대로 HEADING/PARAGRAPH/TABLE 을 쓴다 — "하류는 type 으로만
    분기"(model.py 원칙) 를 지키면서 위치 정보만 추가 메타데이터로 얹는 방식.
  - **이미지**: 인라인 드로잉(w:drawing/a:blip)과 레거시 VML(w:pict/v:imagedata)을 모두
    감지해 FIGURE 블록을 만든다. 본문 단락 이미지는 최상위 블록으로, **표 셀 이미지는
    `TableData.images`([{"row","col","figure"}])** 로 담아 어느 셀에 있던 이미지인지를
    유지한다(pdf/hwpx/hwp 백엔드와 동일 관례). OCR_HOOK 이 등록돼 있으면(doc_parser.register_ocr(),
    PDF 쪽과 동일한 훅) 이미지 바이트를 관계(rId)로 찾아 즉시 OCR 을 돌려 인식된 텍스트를
    채운다 — 등록 안 돼 있으면 text=None 으로 미인식 상태만 남긴다(PDF FIGURE 블록과
    동일 관례). OCR 은 이미지 안 문자만 읽어낼 뿐 "무엇이 그려져 있는지"(장면 의미)는
    여전히 모르므로, OCR 성공 후에도 needs_semantic=True 는 유지한다 — VLM 연결은 별도 TODO.
  - **페이지**: DOCX 에는 고정 페이지 개념이 없다(실제 페이지 수는 렌더링 시점의
    폰트/여백에 따라 달라짐). 여기서는 문서에 **명시적으로 기록된 신호**(수동
    개바꿈 `w:br[@type=page]`, 페이지 단위 구역바꿈)만으로 근사치를 센다 — 임의
    추정이 아니라 실제 XML 신호 기반이지만, 본문 흐름에 의한 자동 줄바꿈은 반영하지
    못하므로 하한값에 가깝다(meta.pages_approx, warnings 에 명시).
"""
from __future__ import annotations

import re
from collections.abc import Callable
from pathlib import Path

import docx
from docx.oxml.ns import qn
from docx.table import Table as DocxTable
from docx.table import _Cell as DocxCell
from docx.text.paragraph import Paragraph as DocxParagraph

from . import ocr_batch
from .model import (
    FIGURE,
    HEADING,
    ORIGIN_TEXT,
    PARAGRAPH,
    SECTION_BODY,
    SECTION_FOOTER,
    SECTION_FOOTER_EVEN,
    SECTION_FOOTER_FIRST,
    SECTION_HEADER,
    SECTION_HEADER_EVEN,
    SECTION_HEADER_FIRST,
    TABLE,
    Block,
    DocumentModel,
    TableData,
)

# 이미지 OCR 훅 — pdf_backend.OCR_HOOK 과 동일 시그니처((이미지 bytes, 미사용 인덱스) ->
# [{"bbox":...,"text":...}])라 doc_parser.register_ocr() 이 등록한 같은 훅 함수(보통
# ocr_paddle.make_ocr_lines_hook())를 그대로 재사용한다. 페이지 개념이 없는 DOCX 이미지엔
# 인덱스 인자가 의미 없어 항상 0 을 넘긴다.
OCR_HOOK: Callable[[bytes, int], list[dict]] | None = None

# mc(Markup Compatibility) 는 python-docx 의 표준 nsmap 에 없어 qn() 이 KeyError 를 낸다.
_MC_FALLBACK = "{http://schemas.openxmlformats.org/markup-compatibility/2006}Fallback"

_HEADING_STYLE_EN = re.compile(r"^Heading\s*(\d+)$", re.IGNORECASE)
_HEADING_STYLE_KO = re.compile(r"^제목\s*(\d+)$")


# ---------------------------------------------------------------------------
# 내용 컨트롤(w:sdt)·필드(w:fldSimple) 정규화 — 파싱 전에 껍데기를 벗긴다
# ---------------------------------------------------------------------------
def _unwrap_shell(root, tag: str, content_tag: str | None = None) -> int:
    """`tag` 엘리먼트를 껍데기만 벗겨 그 내용을 부모의 같은 자리에 그대로 붙인다.
    `content_tag` 가 있으면 그 자식 엘리먼트 안쪽이 실제 내용(w:sdt -> w:sdtContent)."""
    count = 0
    for shell in list(root.iter(tag)):
        parent = shell.getparent()
        if parent is None:      # 바깥 껍데기를 벗기면서 이미 옮겨진 경우
            continue
        content = shell if content_tag is None else shell.find(content_tag)
        idx = parent.index(shell)
        if content is not None:
            for child in list(content):
                parent.insert(idx, child)
                idx += 1
        parent.remove(shell)
        count += 1
    return count


def _normalize_content_controls(root) -> None:
    """`w:sdt`(내용 컨트롤)와 `w:fldSimple`(필드) 껍데기를 벗겨 python-docx 가 원래 보던
    모양으로 되돌린다. python-docx 는 `w:tr`/`w:tc`/`w:p` 의 **직계** 자식만 보기 때문에,
    이 껍데기 안에 든 셀·run 은 존재 자체가 안 보인다.

    실측(2026-08-06, testset docx 10건): 이 조직 템플릿은 입력 필드를 전부 내용 컨트롤로
    만들어 놨다 — `w:tr/w:sdt/w:sdtContent/w:tc` 형태로 **셀 하나가 통째로** sdt 안에 든
    경우 17건, `w:p/w:sdt/w:sdtContent/w:r` 로 문단 일부 run 만 든 경우 6건. 그 결과
    시험의뢰서의 의뢰기관명·대표자·주소·담당자·연락처·이메일·제품명, 을지/갑지의
    의뢰번호·성적서번호가 전부 빈 문자열로 나오고 있었다(표 그리드에서는 `None` 패딩 →
    빈 칸이라 유실인지조차 안 보였다). `w:fldSimple` 은 계산된 결과 run 을 자기 안에
    품고 있어 같은 이유로 꼬리말 "페이지 ( 1 ) / 총 ( 1 )" 이 통째로 유실됐다.

    파싱 로직을 건드리지 않고 입력 XML 만 정규화하는 이유: 표 그리드/병합/중첩 처리는
    python-docx 의 `row.cells`·`grid_cols_before` 에 의존하는데, 여기서 sdt 를 벗겨두면
    그 계산이 원래 의도대로 동작한다(호출부마다 sdt 를 따로 뒤지면 그리드 계산은 여전히
    셀을 못 본다)."""
    _unwrap_shell(root, qn("w:sdt"), qn("w:sdtContent"))
    _unwrap_shell(root, qn("w:fldSimple"))


# ---------------------------------------------------------------------------
# 제목(heading) 판별
# ---------------------------------------------------------------------------
def _pPr_outline_level(pPr) -> int | None:
    """`w:pPr` 엘리먼트에서 `w:outlineLvl`(0-8=제목, 9 또는 없음=본문) 원시값을 읽는다.
    단락(`w:p/w:pPr`)과 스타일 정의(`w:style/w:pPr`) 양쪽에서 재사용."""
    if pPr is None:
        return None
    el = pPr.find(qn("w:outlineLvl"))
    if el is None:
        return None
    val = el.get(qn("w:val"))
    if val is None:
        return None
    lvl = int(val)
    return lvl if lvl < 9 else None


def _outline_level(paragraph: DocxParagraph) -> int | None:
    """단락 자신에게 직접 찍힌 outlineLvl(오버라이드)만 본다. 스타일에서 물려받는
    값은 `_style_chain_outline_level()` 이 따로 본다 — Word 는 보통 outlineLvl 을
    스타일 정의에만 두고 각 단락에 복사해두지 않으므로 이것만으로는 대부분 놓친다."""
    return _pPr_outline_level(paragraph._p.pPr)


def _style_chain_outline_level(paragraph: DocxParagraph) -> int | None:
    """단락이 적용받는 스타일(및 그 부모 스타일 체인)에 정의된 outlineLvl 을 찾는다.

    실측(SST-K-TP-7-01-02 시험의뢰서 등 전체 testset)에서 확인된 패턴: 이 조직 템플릿의
    커스텀 스타일 '본문 타이틀'은 outlineLvl=1 을 **스타일 정의 자체**(styles.xml)에만
    갖고 있고, 그 스타일을 적용한 개별 단락에는 outlineLvl 이 전혀 찍혀있지 않다
    (Word 의 정상적인 스타일 상속 동작 — 각 단락마다 값을 복사하지 않는다). 단락만 보는
    `_outline_level()` 은 이런 케이스를 전부 놓치므로, 실제 OOXML 상속 규칙대로 스타일 →
    부모 스타일 순서로 outlineLvl 을 찾을 때까지 체인을 타고 올라간다.

    이름이 'Heading N'/'제목 N' 관례를 따르지 않는 커스텀 스타일이라도, **스타일 정의에
    outlineLvl 이 있는 경우에만** 잡힌다 — outlineLvl 도 없는 커스텀 스타일(이 조직
    템플릿의 '문서제목'/'서브 타이틀' 등)은 이름 자체를 매칭하는 수밖에 없는데, 그건
    이 조직 템플릿 어휘에 맞춘 튜닝이라 여기서는 하지 않는다(범용 신호만 사용)."""
    style = paragraph.style
    seen: set[int] = set()
    while style is not None and id(style._element) not in seen:
        seen.add(id(style._element))
        lvl = _pPr_outline_level(style._element.pPr)
        if lvl is not None:
            return lvl
        style = style.base_style
    return None


def _heading_level(paragraph: DocxParagraph) -> int | None:
    """제목이면 1-indexed 레벨(Title/Heading 1/제목 1 -> 1)을, 본문이면 None을 반환.
    스타일 이름(영문/국문 관례) 우선 판별 후, 매칭 안 되면 단락 자신의 outlineLvl,
    그마저 없으면 스타일(상속 체인)의 outlineLvl 순으로 폴백한다."""
    style = paragraph.style
    name = (style.name if style is not None else "") or ""
    name = name.strip()
    if name.lower() == "title":
        return 1
    m = _HEADING_STYLE_EN.match(name) or _HEADING_STYLE_KO.match(name)
    if m:
        return int(m.group(1))
    lvl = _outline_level(paragraph)
    if lvl is None:
        lvl = _style_chain_outline_level(paragraph)
    return (lvl + 1) if lvl is not None else None


# ---------------------------------------------------------------------------
# 이미지(인라인 드로잉) 감지 — 내용 인식은 TODO
# ---------------------------------------------------------------------------
def _paragraph_image_rids(paragraph: DocxParagraph) -> list[str]:
    """단락 내 이미지의 관계 ID(r:embed/r:link/r:id) 목록, 중복 제거.

    두 가지 표현 방식을 모두 인식한다 — 실측(SST-K-TI-03-04 을지 문서)에서 이미지 7개 중
    2개만 최신 `w:drawing`(a:blip)이고 나머지 5개는 구버전 호환 모드에서 저장된 레거시
    VML(`w:pict`/`v:imagedata`)이라 하나만 보면 대부분을 놓친다. VML 은 표준 OOXML
    네임스페이스(docx.oxml.ns.nsmap)에 `v` 프리픽스가 없어 local-name() 매칭으로 찾는다.
    `w:drawing` 의 mc:Fallback 으로 같은 이미지의 `w:pict` 표현이 함께 딸려오는 경우가
    있어(Alternate Content), 같은 rId가 두 경로에서 잡히면 중복 FIGURE 블록이 생기지
    않도록 set 으로 걸러낸다. 이미지 자체 내용 인식은 아직 미구현(TODO) — 위치·존재만
    FIGURE 블록으로 남겨 추후 OCR/VLM 연결 지점을 표시."""
    seen: set[str] = set()
    rids: list[str] = []

    def _add(rid: str | None) -> None:
        if rid and rid not in seen:
            seen.add(rid)
            rids.append(rid)

    # 텍스트박스 안 이미지는 제외한다 — 그쪽은 _textbox_blocks() 가 내부 문단을 따로
    # 순회하면서 잡으므로, 여기서도 세면 같은 이미지가 FIGURE 블록 두 개로 나온다.
    for blip in paragraph._p.xpath(".//w:drawing//a:blip[not(ancestor::w:txbxContent)]"):
        _add(blip.get(qn("r:embed")) or blip.get(qn("r:link")))
    for imagedata in paragraph._p.xpath(
            ".//*[local-name()='imagedata'][not(ancestor::*[local-name()='txbxContent'])]"):
        _add(imagedata.get(qn("r:id")) or imagedata.get(qn("r:href")))
    return rids


def _crop_by_rid(paragraph: DocxParagraph) -> dict[str, tuple[float, float, float, float]]:
    """rId → 워드의 **자르기** 비율 (왼쪽, 위, 오른쪽, 아래).

    워드에서 그림을 자르면 원본은 파일에 그대로 남고 `<a:srcRect>` 가 "이만큼은 안
    보인다"고만 적는다. 값은 십만분율이다 — `b="77151"` 이면 아래에서 77.151% 를
    잘라내고 위 22.849% 만 인쇄된다.

    **이걸 안 보면 안 보이는 글자를 문서 내용으로 싣는다.** 실측(제출물 확인증):
    머릿말 레터헤드가 원본 1190×224 인데 위 51px(파란 띠)만 쓴다. 잘린 아래쪽에
    대표·등록번호·전화·주소가 있고, OCR 이 그것을 읽어 머릿말 내용으로 실었다.
    검토자가 문서를 아무리 봐도 없는 문장이 지적의 근거로 떴다.

    한 rId 를 서로 다르게 잘라 여러 번 쓰면 첫 번째 것만 남는다 —
    `_paragraph_image_rids` 가 rId 로 중복을 접기 때문에 짝을 맞춘다.
    """
    out: dict[str, tuple[float, float, float, float]] = {}
    # `pic:` 프리픽스는 python-docx 기본 nsmap 에 없다. VML 쪽과 같이 local-name()
    # 으로 찾는다 — 파일 안에서 한 가지 방식으로 통일한다.
    for fill in paragraph._p.xpath(
            ".//w:drawing//*[local-name()='blipFill']"
            "[not(ancestor::w:txbxContent)]"):
        blips = fill.findall(qn("a:blip"))
        rid = blips[0].get(qn("r:embed")) or blips[0].get(qn("r:link")) if blips else None
        if not rid or rid in out:
            continue
        rect = fill.find(qn("a:srcRect"))
        if rect is None:
            continue
        box = tuple(int(rect.get(k) or 0) / 100000 for k in ("l", "t", "r", "b"))
        if any(box):
            out[rid] = box
    return out


def _apply_crop(image_bytes: bytes, box: tuple[float, float, float, float],
                warnings: list[str]) -> bytes:
    """`srcRect` 가 가리키는 **보이는 부분만** 남긴 이미지.

    자르지 못하면 원본을 그대로 돌려주되 그 사실을 남긴다 — 조용히 원본을 읽으면
    안 보이는 글자가 다시 문서 내용이 된다.
    """
    left, top, right, bottom = box
    try:
        from io import BytesIO

        from PIL import Image
        im = Image.open(BytesIO(image_bytes))
        w, h = im.size
        x0, y0 = round(w * left), round(h * top)
        x1, y1 = round(w * (1 - right)), round(h * (1 - bottom))
        if x1 - x0 < 1 or y1 - y0 < 1:
            # 보이는 영역이 없다시피 하다. 읽을 것이 없으니 OCR 을 걸지 않는다.
            return b""
        buf = BytesIO()
        im.crop((x0, y0, x1, y1)).save(buf, format="PNG")
        return buf.getvalue()
    except Exception as exc:  # noqa: BLE001 — 자르기 실패로 그림을 통째로 버리지 않는다
        warnings.append(
            f"그림의 잘라내기를 적용하지 못해 원본 전체를 읽었습니다({exc}) — "
            "문서에 보이지 않는 글자가 섞일 수 있습니다.")
        return image_bytes


# ---------------------------------------------------------------------------
# 텍스트박스(도형 안 텍스트) — 문단 안에 문단/표가 중첩된 구조
# ---------------------------------------------------------------------------
def _textbox_roots(paragraph: DocxParagraph) -> list:
    """이 문단이 직접 품은 텍스트박스 내용 루트(`w:txbxContent`) 목록.

    텍스트박스는 `w:p/w:r/w:drawing/.../wps:txbx/w:txbxContent/(w:p|w:tbl)` 처럼 **문단
    안에 문단·표가 다시 들어있는** 구조라 `paragraph.text` 로는 한 글자도 안 나온다.
    최신 DrawingML(wps:txbx)과 레거시 VML(v:textbox)이 같은 `w:txbxContent` 를 쓰므로
    태그 하나로 둘 다 잡힌다. 두 가지를 걸러낸다:

      - `mc:Fallback` 안쪽 — Word 는 같은 도형을 mc:Choice(최신)/mc:Fallback(VML) 두 벌로
        싣는다(실측: 갑지 문서의 'N/A' 텍스트박스가 정확히 이 형태). Fallback 을 안 거르면
        같은 내용이 두 번 나온다. Choice 만 읽으면 되고, AlternateContent 는 Choice 없이
        Fallback 만 있는 경우가 없다.
      - 다른 `w:txbxContent` 안쪽 — 중첩 텍스트박스는 바깥에서 한 번, 안쪽 문단을 재귀
        순회할 때 또 한 번 잡히므로 여기서는 가장 바깥 것만 돌려준다."""
    roots = []
    for el in paragraph._p.iter(qn("w:txbxContent")):
        ancestors = {a.tag for a in el.iterancestors()}
        if _MC_FALLBACK in ancestors or qn("w:txbxContent") in ancestors:
            continue
        roots.append(el)
    return roots


def _textbox_blocks(paragraph: DocxParagraph, page: int, section: str,
                    warnings: list[str]) -> list[Block]:
    """텍스트박스 안 문단·표를 본문과 **똑같은 경로**로 변환해 블록 목록으로 돌려준다.

    안쪽 문단은 `_convert_paragraph()`, 안쪽 표는 `_table_blocks()` 를 그대로 재사용하므로
    제목 판별·병합/중첩 표 복원·셀 이미지 수집이 본문과 동일하게 동작하고, 텍스트박스
    안에 또 텍스트박스가 있어도 재귀로 따라간다. 위치는 hwpx/hwp 백엔드가 도형 안 텍스트를
    다루는 방식과 맞춰 문서 최상위 블록으로 흘려보낸다(`section` 은 호스트 문단의 것을 승계)."""
    blocks: list[Block] = []
    for content in _textbox_roots(paragraph):
        for child in content:
            if child.tag == qn("w:p"):
                blocks.extend(_convert_paragraph(DocxParagraph(child, paragraph),
                                                 page, section, warnings))
            elif child.tag == qn("w:tbl"):
                blocks.extend(_table_blocks(DocxTable(child, paragraph),
                                            page, section, warnings))
    return blocks


def _resolve_image_bytes(part, rid: str) -> bytes | None:
    """관계 ID(rId)로 실제 이미지 바이트를 찾는다. part 는 이미지가 속한 스토리 파트
    (document part 또는 header/footer part) — 각 파트가 자기 rels 를 따로 가지므로
    반드시 이미지가 발견된 단락의 `paragraph.part` 를 넘겨야 한다."""
    try:
        return part.rels[rid].target_part.blob
    except Exception:
        return None


def _build_figure_block(part, rid: str, page: int, section: str,
                        warnings: list[str],
                        crop: tuple[float, float, float, float] | None = None) -> Block:
    """이미지 하나를 FIGURE 블록으로 변환. OCR_HOOK 이 등록돼 있으면 이미지 바이트를
    찾아 OCR 을 **예약**한다 — 여기서 기다리지 않고 parse_docx 끝에서 한꺼번에 병렬로
    돌려 이 블록의 text 를 채운다(ocr_batch). 이미지마다 원격 VL 응답(약 5초)을
    차례로 기다리면 그 시간이 그대로 쌓인다.

    PDF FIGURE 블록 관례(text=인식된 내용, 없으면 None)를 그대로 따르고, 관계 해석
    실패는 조용히 넘기지 않고 warnings 에 남긴 뒤 text=None 으로 둔다."""
    block = Block(FIGURE, page, text=None, origin=ORIGIN_TEXT,
                  needs_semantic=True, section=section)
    if OCR_HOOK is not None:
        image_bytes = _resolve_image_bytes(part, rid)
        if image_bytes is None:
            warnings.append(
                f"{ocr_batch.figure_label(page)}을 문서에서 꺼내지 못했습니다. "
                f"이 그림 안의 글자는 검토하지 않았습니다.")
        else:
            # 워드의 **자르기**를 반영한다. 안 하면 잘려서 안 보이는 글자까지 읽어
            # 문서 내용으로 싣는다(_crop_by_rid 참고).
            if crop is not None:
                image_bytes = _apply_crop(image_bytes, crop, warnings)
            if image_bytes:
                ocr_batch.schedule(block, image_bytes, warnings)
    return block


# ---------------------------------------------------------------------------
# 페이지 근사 — 명시적 개바꿈/구역바꿈 신호만 반영
# ---------------------------------------------------------------------------
def _has_manual_page_break(paragraph: DocxParagraph) -> bool:
    return bool(paragraph._p.xpath('.//w:br[@w:type="page"]'))


def _starts_new_page_section(paragraph: DocxParagraph) -> bool:
    """이 단락이 구역(section)의 끝이면서, 새 구역이 새 페이지에서 시작하는 유형
    (nextPage/oddPage/evenPage — continuous 가 아님)인지. `w:sectPr` 는 구역의 마지막
    단락의 `w:pPr` 안에 온다(OOXML 관례)."""
    pPr = paragraph._p.pPr
    if pPr is None:
        return False
    sectPr = pPr.find(qn("w:sectPr"))
    if sectPr is None:
        return False
    type_el = sectPr.find(qn("w:type"))
    if type_el is None:
        return True  # 기본값 nextPage
    val = type_el.get(qn("w:val"))
    return val in (None, "nextPage", "oddPage", "evenPage")


# ---------------------------------------------------------------------------
# 표 — 그리드/병합/중첩을 python-docx 에서 직접 복원
# ---------------------------------------------------------------------------
def _table_grid(table: DocxTable) -> tuple[int, list[list[DocxCell | None]]]:
    """`table.rows[i].cells` (병합 셀은 앵커 셀 참조가 반복됨)를 grid_cols_before/after
    로 보정해 완전한 n_rows x n_cols 그리드로 만든다. Word 는 표가 첫/끝 열에서
    "일찍 시작/끝나는" 행을 허용하므로 그 자리는 None 으로 채운다."""
    n_cols = len(table.columns)
    grid: list[list[DocxCell | None]] = []
    for row in table.rows:
        slots: list[DocxCell | None] = [None] * row.grid_cols_before
        slots.extend(row.cells)
        slots.extend([None] * row.grid_cols_after)
        if len(slots) < n_cols:
            slots.extend([None] * (n_cols - len(slots)))
        grid.append(slots[:n_cols])
    return n_cols, grid


def _extract_table(table: DocxTable, page: int, section: str,
                   warnings: list[str], extra_blocks: list[Block]) -> TableData:
    """표 하나를 TableData 로 변환. 같은 병합 셀은 python-docx 가 동일한 `w:tc` 엘리먼트를
    그리드 여러 칸에서 참조하므로, `id(cell._tc)` 로 병합 영역을 정확히 식별한다
    (근사·휴리스틱이 아니라 OOXML 구조 자체에서 나오는 값). 병합 영역은 좌상단(앵커)
    칸에만 텍스트를 채우고 나머지는 ""로 비워 pdf_backend/docling_adapter 관례와 맞춘다.
    셀 안에 중첩된 표(cell.tables)는 재귀적으로 같은 방식으로 변환한다.

    셀 안 이미지는 `TableData.images`([{"row","col","figure": Block}])에 담는다 — 실측
    (SST-K-TI-03-04 을지 문서)에서 이미지 7개가 전부 표 셀 안에 있었는데, 셀 텍스트
    (cell.text)만 보고 지나가면 이미지 존재 자체가 통째로 유실된다. 표/셀 단락도 본문
    단락과 동일하게 `_paragraph_image_rids` 로 훑는다. 중첩 표 안의 이미지는 그 중첩
    TableData 자신의 images 에 담겨 (row,col) 이 어느 표 기준인지 섞이지 않는다.
    pdf/hwpx/hwp 백엔드와 동일한 관례 — 예전처럼 표 뒤에 형제 FIGURE 블록으로 흘려보내면
    "어느 셀에 있던 이미지인지"가 유실되고, table.images 만 보는 하류에는 DOCX 셀
    이미지만 통째로 안 보인다."""
    n_cols, grid = _table_grid(table)
    n_rows = len(grid)
    cells_text: list[list[str]] = [["" for _ in range(n_cols)] for _ in range(n_rows)]

    positions: dict[int, list[tuple[int, int]]] = {}
    cell_by_key: dict[int, DocxCell] = {}
    for r in range(n_rows):
        for c in range(n_cols):
            cell = grid[r][c]
            if cell is None:
                continue
            key = id(cell._tc)
            positions.setdefault(key, []).append((r, c))
            cell_by_key[key] = cell

    merges: list[dict[str, int]] = []
    nested_tables: list[dict] = []
    images: list[dict] = []
    for key, pos_list in positions.items():
        rs = [p[0] for p in pos_list]
        cs = [p[1] for p in pos_list]
        anchor_r, anchor_c = min(rs), min(cs)
        row_span, col_span = max(rs) - anchor_r + 1, max(cs) - anchor_c + 1
        cell = cell_by_key[key]
        cells_text[anchor_r][anchor_c] = cell.text.strip()
        if row_span > 1 or col_span > 1:
            merges.append({"row": anchor_r, "col": anchor_c,
                           "row_span": row_span, "col_span": col_span})
        for paragraph in cell.paragraphs:
            # part 를 단락에서 가져오는 이유: 관계(rId)는 파트마다 독립적이라 이미지가
            # 발견된 단락 자신의 part 로만 바이트를 되찾을 수 있다(머리말/꼬리말 표).
            crops = _crop_by_rid(paragraph)
            for rid in _paragraph_image_rids(paragraph):
                images.append({
                    "row": anchor_r, "col": anchor_c,
                    "figure": _build_figure_block(paragraph.part, rid, page, section,
                                                  warnings, crops.get(rid)),
                })
            # 셀 안 텍스트박스의 내용은 표 격자에 넣을 자리가 없어(문단·표가 통째로 들어있다)
            # 문서 최상위 블록으로 흘려보낸다 — hwpx 백엔드의 도형 텍스트 처리와 같은 관례.
            extra_blocks.extend(_textbox_blocks(paragraph, page, section, warnings))
        for nested in cell.tables:
            nested_tables.append({
                "row": anchor_r, "col": anchor_c,
                "table": _extract_table(nested, page, section, warnings, extra_blocks),
            })

    return TableData(rows=n_rows, cols=n_cols, cells=cells_text,
                     nested=bool(nested_tables), detected_only=False,
                     merges=merges, nested_tables=nested_tables, images=images)


def _table_blocks(table: DocxTable, page: int, section: str,
                  warnings: list[str]) -> list[Block]:
    """표 블록 + 그 표(중첩 표 포함) 셀 안 텍스트박스에서 나온 블록들.
    셀 안 이미지는 형제 블록이 아니라 `TableData.images` 에 (row,col) 과 함께 들어간다
    (_extract_table 참조)."""
    extra_blocks: list[Block] = []
    table_data = _extract_table(table, page, section, warnings, extra_blocks)
    return [Block(TABLE, page, origin=ORIGIN_TEXT, section=section, table=table_data),
            *extra_blocks]


# ---------------------------------------------------------------------------
# 단락 -> 블록
# ---------------------------------------------------------------------------
def _convert_paragraph(paragraph: DocxParagraph, page: int, section: str,
                       warnings: list[str]) -> list[Block]:
    blocks: list[Block] = []
    text = paragraph.text.strip()
    if text:
        level = _heading_level(paragraph)
        if level is not None:
            blocks.append(Block(HEADING, page, text=text, level=level,
                                origin=ORIGIN_TEXT, section=section))
        else:
            blocks.append(Block(PARAGRAPH, page, text=text,
                                origin=ORIGIN_TEXT, section=section))
    crops = _crop_by_rid(paragraph)
    for rid in _paragraph_image_rids(paragraph):
        blocks.append(_build_figure_block(paragraph.part, rid, page, section,
                                          warnings, crops.get(rid)))
    blocks.extend(_textbox_blocks(paragraph, page, section, warnings))
    return blocks


# ---------------------------------------------------------------------------
# 머리말/꼬리말
# ---------------------------------------------------------------------------
def _header_footer_blocks(document, warnings: list[str]) -> list[Block]:
    """구역(section)마다 머리말/꼬리말을 Block.section 값으로 종류를 구분해 뽑는다.

    기본(primary) 머리말/꼬리말 외에 두 가지 특수 정의도 다룬다 — 단, **Word 가 실제로
    그 정의를 쓰는 조건이 켜져 있을 때만** 블록을 만든다(꺼져 있으면 XML 에 정의가
    남아있어도 Word 는 렌더링에서 무시한다, python-docx 자체 문서에 명시):
      - 첫 페이지 전용: `section.different_first_page_header_footer` 가 True 인
        구역에서만 `first_page_header`/`first_page_footer` 를 읽는다.
      - 짝수 페이지 전용: 문서 설정 `document.settings.odd_and_even_pages_header_footer`
        가 True 일 때만 `even_page_header`/`even_page_footer` 를 읽는다(구역이 아니라
        문서 전체 설정).

    각 정의는 이전 구역과 같은 걸 상속(is_linked_to_previous=True)하면 건너뛰어
    중복을 막고, 같은 정의를 참조하는 경우도 파트 identity(id(hf.part))로 한 번만
    처리한다 — hf.part 는 python-docx 가 관계(rId)마다 캐싱해 반환하므로 같은 정의는
    항상 같은 객체가 나온다(반복 호출로 확인됨)."""
    blocks: list[Block] = []
    seen_part_ids: set[int] = set()
    use_even = document.settings.odd_and_even_pages_header_footer

    def _emit(hf, section_tag: str) -> None:
        if hf.is_linked_to_previous:
            return
        part_id = id(hf.part)
        if part_id in seen_part_ids:
            return
        seen_part_ids.add(part_id)
        # 머리말/꼬리말은 각자 별도 XML 파트라 본문과 따로 정규화해야 한다(꼬리말의
        # w:fldSimple 페이지 필드가 여기서 살아난다). 파트당 한 번만 — 위 dedupe 덕분.
        _normalize_content_controls(hf.part.element)
        for item in hf.iter_inner_content():
            if isinstance(item, DocxTable):
                blocks.extend(_table_blocks(item, 0, section_tag, warnings))
            else:
                blocks.extend(_convert_paragraph(item, 0, section_tag, warnings))

    for sec in document.sections:
        _emit(sec.header, SECTION_HEADER)
        _emit(sec.footer, SECTION_FOOTER)
        if sec.different_first_page_header_footer:
            _emit(sec.first_page_header, SECTION_HEADER_FIRST)
            _emit(sec.first_page_footer, SECTION_FOOTER_FIRST)
        if use_even:
            _emit(sec.even_page_header, SECTION_HEADER_EVEN)
            _emit(sec.even_page_footer, SECTION_FOOTER_EVEN)
    return blocks


# ---------------------------------------------------------------------------
# 메인
# ---------------------------------------------------------------------------
def parse_docx(path: str | Path) -> DocumentModel:
    """doc_parser.parse_document() 가 .docx 확장자에 대해 위임하는 최종 구현.
    OOXML 을 python-docx 로 직접 순회해 표(병합/중첩 포함)·제목·머리말/꼬리말·이미지
    위치를 공통 문서 모델로 정규화한다. PDF 백엔드와 달리 판별/OCR 단계가 없다."""
    path = str(path)
    name = Path(path).name
    warnings: list[str] = []
    try:
        document = docx.Document(path)
    except Exception as e:  # noqa
        warnings.append(f"docx 열기 실패: {e}")
        return DocumentModel(name, {"opened": False}, [], warnings)

    _normalize_content_controls(document.element)
    # 순회 중에 만나는 이미지는 OCR 을 예약만 하고(_build_figure_block), 순회가 끝난
    # 뒤 아래에서 한꺼번에 병렬로 돌린다. begin() 은 앞선 파싱이 중간에 터져 남긴
    # 예약이 이 문서로 새지 않게 비운다.
    ocr_batch.begin()
    blocks: list[Block] = _header_footer_blocks(document, warnings)

    page = 0
    for item in document.iter_inner_content():
        if isinstance(item, DocxTable):
            blocks.extend(_table_blocks(item, page, SECTION_BODY, warnings))
            continue
        blocks.extend(_convert_paragraph(item, page, SECTION_BODY, warnings))
        if _has_manual_page_break(item):
            page += 1
        if _starts_new_page_section(item):
            page += 1

    meta = {
        "opened": True,
        "paragraphs": len(document.paragraphs),
        "tables": len(document.tables),
        "sections": len(document.sections),
        "pages_approx": page + 1,
    }
    # 검토자 화면에 지적으로 나가면 안 되는 종류다 — 이 문서를 못 읽었다는
    # 보고가 아니라 **우리 숫자에 붙는 각주**다. warnings 에 넣으면
    # parser_bridge 가 그대로 parser_warnings 로 옮기고 orchestrator 가
    # 미검토 INFO 카드로 만든다(모든 문서에서 매번).
    meta["backend_notes"] = [
        "pages_approx 는 명시적 개바꿈/구역바꿈만 반영한 근사 하한값 — "
        "실제 워드 렌더링 페이지수(본문 흐름에 의한 자동 줄바꿈 포함)와 다를 수 있음"]
    ocr_batch.run(OCR_HOOK)
    return DocumentModel(name, meta, blocks, warnings)
